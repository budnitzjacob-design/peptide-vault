"""Generate Peptide Vault implementation plan as Word doc."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Styles ─────────────────────────────────────────────────────────────────
styles = doc.styles

def add_heading(text, level=1, color=(0, 82, 155)):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.color.rgb = RGBColor(*color)
    return h

def add_body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for para in hdr_cells[i].paragraphs:
            for run in para.runs:
                run.bold = True
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_text in enumerate(row_data):
            row_cells[i].text = cell_text
    doc.add_paragraph()
    return table

def add_code(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.left_indent = Inches(0.3)
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
    return p

# ── Title Page ──────────────────────────────────────────────────────────────
title = doc.add_heading('PEPTIDE VAULT', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(0, 82, 155)
    run.font.size = Pt(28)

sub = doc.add_paragraph('Implementation Plan & Technical Specification')
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.size = Pt(14)
sub.runs[0].font.italic = True

meta = doc.add_paragraph(f'Author: Jacob Budnitz  |  Date: April 20, 2026  |  Version: 1.0')
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.runs[0].font.size = Pt(10)

doc.add_paragraph()
doc.add_paragraph('Confidential – Internal Use Only').alignment == WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# ── Table of Contents ────────────────────────────────────────────────────────
add_heading('Table of Contents', 1)
toc_items = [
    '1. Executive Summary',
    '2. Project Overview & Goals',
    '3. Site Architecture',
    '4. Data Model & Schema',
    '5. Feature Specifications',
    '   5.1  Filtering & Sorting',
    '   5.2  Peptide Cards (Collapsed / Expanded)',
    '   5.3  Protein Cascade Accordions',
    '   5.4  User Authentication & Bookmarks',
    '   5.5  Search',
    '6. Design System',
    '7. Data Pipeline & Scraping Strategy',
    '8. Technology Stack',
    '9. File & Directory Structure',
    '10. Development Phases & Timeline',
    '11. Deployment (GitHub + Fly.io)',
    '12. Future Enhancements',
    '13. Open Questions',
]
for item in toc_items:
    add_bullet(item)

doc.add_page_break()

# ── 1. Executive Summary ─────────────────────────────────────────────────────
add_heading('1. Executive Summary', 1)
add_body(
    'Peptide Vault is a scientifically rigorous, publicly accessible web application '
    'for browsing, filtering, and researching therapeutic peptides. Modeled after '
    'peptpedia.org but substantially expanded, it provides structured data on amino '
    'acid sequences, pharmacokinetics, protein cascades, clinical effects, safety '
    'profiles, and peer-reviewed citations for 41+ peptides.'
)
add_body(
    'The site targets researchers, clinicians, biohackers, and students. Users can '
    'create accounts to bookmark compounds. All data is cited, hyperlinked, and '
    'organized with expandable detail sections. The build progresses from an HTML '
    'prototype to a full-stack Node.js application deployed on Fly.io.'
)

doc.add_page_break()

# ── 2. Project Overview ──────────────────────────────────────────────────────
add_heading('2. Project Overview & Goals', 1)
add_heading('Primary Goals', 2, color=(30, 120, 180))
goals = [
    'Comprehensive peptide database with 41+ entries from peptpedia + additional research sources',
    'Multi-axis filtering: chemical family, clinical effects (subcategorized), safety rating',
    'Expandable tiles with amino acid sequence, half-life, dosing, protein cascades, citations',
    'Human use / animal model / performance enhancement summaries per peptide',
    'User authentication (JWT) + bookmark system',
    'Scientific visual design consistent with organism-logger.fly.dev and jacobbudnitz.com',
    'All citations hyperlinked to PubMed / DOI',
    'Deployed on Fly.io with GitHub CI/CD',
]
for g in goals:
    add_bullet(g)

add_heading('Non-Goals (v1)', 2, color=(30, 120, 180))
non_goals = [
    'Mobile app (web responsive only)',
    'Peptide comparison tool (future)',
    'User-submitted data / community edits (future)',
    'Paid tier or paywalled content',
]
for ng in non_goals:
    add_bullet(ng)

doc.add_page_break()

# ── 3. Site Architecture ─────────────────────────────────────────────────────
add_heading('3. Site Architecture', 1)

add_heading('Phase 1 — HTML Prototype', 2, color=(30, 120, 180))
add_body(
    'A single self-contained index.html with embedded CSS and JavaScript. '
    'All 41 peptides stored in an inline JSON object or imported data.js file. '
    'Full UI functionality including filtering, sorting, search, expanded tiles, '
    'and mock auth via localStorage. No server required — opens directly in browser.'
)

add_heading('Phase 2 — Full-Stack App', 2, color=(30, 120, 180))
add_body(
    'Node.js/Express REST API serving peptide data from a structured JSON/SQLite store. '
    'PostgreSQL (Fly.io managed Postgres) for user accounts and bookmarks. '
    'JWT-based authentication. React or vanilla JS SPA frontend.'
)

add_heading('Phase 3 — Production Deploy', 2, color=(30, 120, 180))
add_body(
    'GitHub repository with Actions CI (lint, test, build). '
    'Dockerized Node.js app deployed to Fly.io. '
    'Fly.io managed Postgres for persistence. '
    'Static assets served via Fly.io CDN or optional Cloudflare.'
)

add_heading('System Diagram (Text)', 2, color=(30, 120, 180))
diagram_lines = [
    'Browser',
    '  └─ index.html / SPA (HTML + CSS + Vanilla JS / React)',
    '       ├─ GET /api/peptides        → Node.js/Express',
    '       │      └─ peptides.json / PostgreSQL',
    '       ├─ POST /api/auth/register  → bcrypt + JWT',
    '       ├─ POST /api/auth/login     → JWT response',
    '       └─ GET/POST /api/bookmarks  → PostgreSQL (users, bookmarks)',
    '',
    'Fly.io',
    '  ├─ peptide-vault app (Node.js, Dockerfile)',
    '  └─ Fly.io Postgres (managed)',
    '',
    'GitHub',
    '  ├─ Source code repo',
    '  └─ Actions CI → fly deploy on main merge',
]
for line in diagram_lines:
    add_code(line)

doc.add_page_break()

# ── 4. Data Model ────────────────────────────────────────────────────────────
add_heading('4. Data Model & Schema', 1)

add_heading('4.1 Peptide Object (JSON)', 2, color=(30, 120, 180))
schema_lines = [
    '{',
    '  "id": "bpc-157",',
    '  "slug": "bpc-157",',
    '  "names": {',
    '    "trade": "BPC-157",',
    '    "scientific": "Body Protection Compound-157",',
    '    "proper": "Pentadecapeptide BPC 157",',
    '    "aliases": ["BPC157", "PL 14736"]',
    '  },',
    '  "sequence": {',
    '    "one_letter": "GEPPPGKPADDAGLV",',
    '    "three_letter": "Gly-Glu-Pro-Pro-Pro-Gly-Lys-Pro-Ala-Asp-Asp-Ala-Gly-Leu-Val",',
    '    "length": 15',
    '  },',
    '  "molecular": {',
    '    "formula": "C62H98N16O22",',
    '    "weight": 1419.53,',
    '    "cas": "137525-51-0",',
    '    "chemical_family": "Gastric peptide",',
    '    "structure_class": "Linear peptide"',
    '  },',
    '  "pharmacokinetics": {',
    '    "half_life": "15.2 minutes (IV, rat)",',
    '    "bioavailability": { "SC": "moderate", "oral": "high (gastric stable)" },',
    '    "routes": ["SC", "IM", "IV", "oral", "IP"]',
    '  },',
    '  "dosing": {',
    '    "routes": [',
    '      { "route": "SC", "dose": "10-20 mcg/kg", "frequency": "Once daily" },',
    '      { "route": "oral", "dose": "10 mcg/kg - 10 mg/kg", "frequency": "Once daily" }',
    '    ],',
    '    "notes": "No established human clinical dose"',
    '  },',
    '  "cascades": {',
    '    "metabolism": ["Nitric oxide modulation", "eNOS phosphorylation"],',
    '    "energy_generation": [],',
    '    "growth": ["GHR upregulation", "FAK-Paxillin pathway"],',
    '    "proliferation": ["Fibroblast proliferation", "Angiogenesis via Egr-1/NAB2"],',
    '    "cancer": ["Aberrant angiogenesis prevention"],',
    '    "central_dogma": ["c-Fos/c-Jun/Egr-1 transcription factor activation"],',
    '    "immune_function": ["TNF-alpha↓", "IL-1beta↓", "IL-6↓"],',
    '    "cytokines": ["TNF-alpha↓", "IL-1beta↓", "IL-6↓"]',
    '  },',
    '  "effects": {',
    '    "clinical_categories": ["Recovery & Repair", "GI Protection"],',
    '    "clinical_effects": [',
    '      { "name": "Tissue Repair", "subcategory": "Musculoskeletal", "evidence": "Extensive" },',
    '      { "name": "GI Protection", "subcategory": "Gastrointestinal", "evidence": "Extensive" }',
    '    ],',
    '    "safety_rating": 3,',
    '    "human_use": {',
    '      "summary": "No human clinical trials completed as of 2026...",',
    '      "citations": []',
    '    },',
    '    "animal_models": {',
    '      "summary": ">100 preclinical studies, primarily rat models...",',
    '      "citations": [{ "authors": "Sikiric et al.", "year": 2011, "pmid": "..." }]',
    '    },',
    '    "performance": {',
    '      "summary": "Studied for athletic recovery and injury repair...",',
    '      "citations": []',
    '    }',
    '  },',
    '  "safety": {',
    '    "rating": 3,',
    '    "label": "Preclinical — human data absent",',
    '    "contraindications": [],',
    '    "adverse_events": "None reported in animal studies"',
    '  },',
    '  "citations": [',
    '    {',
    '      "id": "sikiric2011",',
    '      "authors": "Sikiric P, et al.",',
    '      "year": 2011,',
    '      "title": "Stable gastric pentadecapeptide BPC 157...",',
    '      "journal": "Current Pharmaceutical Design",',
    '      "pmid": "21352100",',
    '      "url": "https://pubmed.ncbi.nlm.nih.gov/21352100/"',
    '    }',
    '  ],',
    '  "research_summary": "BPC-157 is a 15-amino-acid synthetic peptide...",',
    '  "last_updated": "2026-04-20"',
    '}',
]
for line in schema_lines:
    add_code(line)

add_heading('4.2 SQL Schema (PostgreSQL)', 2, color=(30, 120, 180))
sql_lines = [
    'CREATE TABLE users (',
    '  id          UUID PRIMARY KEY DEFAULT gen_random_uuid(),',
    '  email       TEXT UNIQUE NOT NULL,',
    '  password_hash TEXT NOT NULL,',
    '  created_at  TIMESTAMPTZ DEFAULT NOW()',
    ');',
    '',
    'CREATE TABLE bookmarks (',
    '  id          UUID PRIMARY KEY DEFAULT gen_random_uuid(),',
    '  user_id     UUID REFERENCES users(id) ON DELETE CASCADE,',
    '  peptide_id  TEXT NOT NULL,',
    '  created_at  TIMESTAMPTZ DEFAULT NOW(),',
    '  UNIQUE(user_id, peptide_id)',
    ');',
]
for line in sql_lines:
    add_code(line)

doc.add_page_break()

# ── 5. Feature Specifications ────────────────────────────────────────────────
add_heading('5. Feature Specifications', 1)

add_heading('5.1 Filtering & Sorting', 2, color=(30, 120, 180))
add_body('Filter bar fixed at top of peptide grid. All filters are combinable. URL query params reflect active filters for shareability.')

add_heading('Filter Axes:', 3, color=(60, 140, 200))
filter_rows = [
    ['Chemical Family / Structure', 'Linear peptide, Cyclic peptide, Glycopeptide, Lipopeptide, GLP-1 analog, GH secretagogue, Melanocortin, Nootropic, Tetrapeptide, Tripeptide'],
    ['Clinical Effect (top-level)', 'Recovery & Repair, Metabolic, Cognitive Enhancement, Anti-Aging, Immune Support, Sexual Medicine, Cardiovascular, GI / Gut Health, Endocrine, Neuroprotection'],
    ['Clinical Effect Subcategory', 'Tissue Repair, Wound Healing, Weight Loss, Glycemic Control, Memory, Neurogenesis, Longevity, Hair, Skin, Libido, Cardiac, Hepatic, etc.'],
    ['Safety Rating', '1 (FDA-approved) → 2 (human trials) → 3 (animal models only) → 4 (limited data) → 5 (theoretical)'],
    ['Half-Life Range', 'Ultra-short (<30 min), Short (30 min–4 h), Medium (4–24 h), Long (1–8 days)'],
    ['Research Volume', 'Extensive (>50 studies), Moderate (10–50), Preliminary (<10)'],
]
add_table(['Filter', 'Options'], filter_rows)

add_heading('Sort Options:', 3, color=(60, 140, 200))
sort_options = [
    'Name (A–Z / Z–A)',
    'Half-life (shortest / longest first)',
    'Safety rating (safest / most experimental first)',
    'Research volume (most to least studied)',
    'Recently updated',
]
for s in sort_options:
    add_bullet(s)

add_heading('5.2 Peptide Cards — Collapsed State', 2, color=(30, 120, 180))
collapsed_fields = [
    ('Top of card', 'Peptide name (trade name large, scientific name small)'),
    ('Name row', 'Proper name + aliases in small text'),
    ('Badge row', 'Chemical family badge | Safety rating indicator (colored dot + label)'),
    ('Sequence', 'Amino acid sequence in monospace font (JetBrains Mono), scrollable if long'),
    ('Stats row', 'Half-life | Molecular weight | AA count'),
    ('Effects pills', 'Top 3 clinical effects as colored pills'),
    ('Expand button', 'Chevron / "Expand" to reveal full detail'),
    ('Bookmark icon', 'Star icon (top-right corner); fills yellow on bookmark'),
]
add_table(['Field', 'Description'], collapsed_fields)

add_heading('5.3 Peptide Cards — Expanded State', 2, color=(30, 120, 180))
expanded_fields = [
    ('Full Names', 'Trade name, scientific name, proper name, all aliases, CAS number'),
    ('Amino Acid Sequence', 'Full sequence with one-letter and three-letter codes; residue count; visual sequence bar'),
    ('Molecular Profile', 'Formula, molecular weight, CAS number, structure class, solubility'),
    ('Pharmacokinetics Table', 'Half-life (by route), bioavailability, peak time, storage'),
    ('Dosing Table', 'Route | Dose | Frequency | Notes — with disclaimer'),
    ('Protein Cascade Accordions', 'See 5.4 below'),
    ('Research Overview', 'Human use summary | Animal models summary | Performance/enhancement summary — each with citations'),
    ('All Clinical Effects', 'Full list with evidence level (Extensive / Moderate / Preliminary)'),
    ('Safety Profile', 'Rating + label, contraindications, adverse events, regulatory status'),
    ('Full Citations', 'Numbered list, hyperlinked to PubMed/DOI'),
    ('Related Peptides', 'Links to similar compounds'),
]
add_table(['Field', 'Description'], expanded_fields)

add_heading('5.4 Protein Cascade Accordions', 2, color=(30, 120, 180))
add_body(
    'Within each expanded tile, protein cascades are displayed as an accordion section. '
    'The top level shows 8 cascade categories as collapsible rows. '
    'Each category shows a one-line summary with element count. '
    'Clicking a category reveals the full mechanism: pathway elements listed with '
    'brief descriptions and citation superscripts.'
)
cascade_categories = [
    ('Metabolism', 'Metabolic pathways, enzyme activations, substrate processing'),
    ('Energy Generation', 'Mitochondrial function, ATP pathways, oxidative phosphorylation'),
    ('Growth', 'GH/IGF-1 axis, bone/muscle growth signals, fibroblast activity'),
    ('Proliferation', 'Cell division, angiogenesis, tissue regeneration pathways'),
    ('Cancer', 'Oncogenic suppression, metastatic gene suppression, tumor models'),
    ('Central Dogma', 'Gene expression, transcription factors, epigenetic regulation'),
    ('Immune Function', 'Innate/adaptive immune modulation, thymic function'),
    ('Cytokines', 'Pro- and anti-inflammatory cytokine effects (IL, TNF, IFN, etc.)'),
]
add_table(['Category', 'What It Covers'], cascade_categories)

add_heading('5.5 User Authentication & Bookmarks', 2, color=(30, 120, 180))
auth_features = [
    'Sign-up modal: email + password (bcrypt hashed, min 8 chars)',
    'Login modal: email + password → returns JWT (stored in localStorage)',
    'JWT decoded on each page load; expired tokens trigger re-login prompt',
    'Bookmark icon on every peptide card; requires login to save',
    'Bookmarks page: filterable/sortable list of saved peptides',
    'Account page: email display, change password, logout, delete account',
    'Rate limiting: 10 login attempts per 15 min per IP',
]
for f in auth_features:
    add_bullet(f)

add_heading('5.6 Search', 2, color=(30, 120, 180))
search_features = [
    'Fuzzy search across: all name fields, CAS number, amino acid sequence fragment',
    'Search bar in header, persistent across pages',
    'Debounced (300 ms) client-side filtering on the JSON dataset',
    'Highlighted matching text in results',
]
for f in search_features:
    add_bullet(f)

doc.add_page_break()

# ── 6. Design System ─────────────────────────────────────────────────────────
add_heading('6. Design System', 1)

add_heading('Color Palette', 2, color=(30, 120, 180))
color_rows = [
    ['Background Primary', '#0a0b0f', 'Near-black — main page bg'],
    ['Background Card', '#111318', 'Slightly lighter — card bg'],
    ['Background Elevated', '#1a1d26', 'Modals, dropdowns'],
    ['Border Subtle', '#2a2d3a', 'Card borders, dividers'],
    ['Accent Cyan', '#00d4c8', 'Primary CTA, links, hover states'],
    ['Accent Blue', '#4f8ef7', 'Secondary accent, badges'],
    ['Text Primary', '#f0f2f5', 'Main content text'],
    ['Text Secondary', '#8b9ab0', 'Labels, meta text'],
    ['Safety — Approved', '#22c55e', 'FDA-approved (rating 1)'],
    ['Safety — Human Trials', '#84cc16', 'Human trial data (rating 2)'],
    ['Safety — Animal Only', '#f59e0b', 'Animal models only (rating 3)'],
    ['Safety — Limited', '#ef4444', 'Limited/no data (ratings 4–5)'],
    ['Sequence BG', '#0d1117', 'Monospace sequence blocks'],
]
add_table(['Token', 'Value', 'Usage'], color_rows)

add_heading('Typography', 2, color=(30, 120, 180))
type_rows = [
    ['Peptide Names (large)', 'Inter, 700, 20px', 'Card headers'],
    ['Scientific Names', 'Inter, 400 italic, 13px', 'Sub-headers'],
    ['Body Text', 'Inter, 400, 14px', 'Descriptions, summaries'],
    ['Monospace (sequences)', 'JetBrains Mono, 400, 13px', 'AA sequences, formulas, code'],
    ['Labels / Badges', 'Inter, 600, 11px uppercase', 'Effect pills, category tags'],
    ['Citations', 'Inter, 400, 12px', 'Reference lists'],
]
add_table(['Element', 'Font Spec', 'Used For'], type_rows)

add_heading('Layout', 2, color=(30, 120, 180))
layout_items = [
    'Responsive grid: 1 col (mobile) → 2 col (tablet) → 3 col (desktop ≥1280px)',
    'Sticky filter/sort bar below fixed header',
    'Cards use CSS Grid with auto-expand animation (max-height transition)',
    'Dark glassmorphism cards: backdrop-filter blur, semi-transparent bg, hairline border',
    'Academic citation numbering (superscript [1], [2]…) with footnote list at bottom of expanded card',
    'Sequence display: scrollable horizontal monospace block with residue numbering every 10 AAs',
]
for item in layout_items:
    add_bullet(item)

add_heading('Scientific Aesthetic Details', 2, color=(30, 120, 180))
aesthetic_items = [
    'Grid overlay texture (subtle) on hero section, like a microarray or cell culture plate',
    'Molecular formula displayed with subscript numbers (e.g., C₆₂H₉₈)',
    'Safety rating displayed as colored dot + text label, not just a number',
    'Evidence level badges: "EXTENSIVE", "MODERATE", "PRELIMINARY" in color-coded all-caps',
    'Citation superscripts clickable — jump to citation block at bottom of card',
    'Small disclaimer ribbon at bottom of each expanded card ("For research purposes only")',
    'Header: monospace logo text "PEPTIDE VAULT" with subtle green-teal glow',
]
for item in aesthetic_items:
    add_bullet(item)

doc.add_page_break()

# ── 7. Data Pipeline ─────────────────────────────────────────────────────────
add_heading('7. Data Pipeline & Scraping Strategy', 1)

add_heading('7.1 Sources', 2, color=(30, 120, 180))
sources = [
    ('peptpedia.org', 'Primary — all 41 peptides; detailed pages per peptide', 'WebFetch per slug'),
    ('PubMed (pubmed.ncbi.nlm.nih.gov)', 'Primary literature; verifies citations from peptpedia', 'WebSearch per peptide name + key terms'),
    ('PubChem (pubchem.ncbi.nlm.nih.gov)', 'Molecular data, CAS numbers, structure images, SMILES', 'WebFetch /compound/<name>/JSON'),
    ('DrugBank (drugbank.com)', 'FDA-approved entries (Semaglutide, PT-141, Pramlintide)', 'WebFetch per entry'),
    ('ClinicalTrials.gov', 'Human trial history and status', 'WebSearch per NCT numbers'),
    ('Preprints.org / bioRxiv', 'Cutting-edge peptide research 2024–2026', 'WebSearch per peptide'),
]
add_table(['Source', 'Purpose', 'Method'], sources)

add_heading('7.2 Scraping Process', 2, color=(30, 120, 180))
add_body('All scraping is done via WebFetch (built-in tool) and WebSearch. No external API keys needed.')
steps = [
    'Step 1: Fetch all 41 peptide pages from peptpedia.org/peptide/<slug>',
    'Step 2: Parse into structured JSON per the data model in Section 4',
    'Step 3: For each peptide, run WebSearch for "[name] amino acid sequence half life mechanisms pubmed" to augment data',
    'Step 4: Fetch PubChem compound data for molecular details',
    'Step 5: Verify citation PMIDs and build hyperlinked citation objects',
    'Step 6: Write compiled data to data/peptides.json and data/scraped/peptpedia_scrape.txt',
    'Step 7: Manual review pass for accuracy; annotate confidence levels',
]
for s in steps:
    add_bullet(s)

add_heading('7.3 Consensus MCP Note', 2, color=(30, 120, 180))
add_body(
    'NOTE: Consensus MCP was not detected in the current tool session. '
    'If connected, it would be used at Step 3 to pull peer-reviewed summaries '
    'directly from Consensus search API, enriching the animal models and human '
    'use summaries with higher citation density. '
    'Workaround: WebSearch targeted at PubMed and preprint servers achieves similar results. '
    'Connect Consensus via Claude.ai integrations if available.'
)

doc.add_page_break()

# ── 8. Technology Stack ──────────────────────────────────────────────────────
add_heading('8. Technology Stack', 1)

tech_rows = [
    ['Frontend (Phase 1)', 'Vanilla HTML/CSS/JS', 'No build step; instant prototype'],
    ['Frontend (Phase 2+)', 'Vanilla JS or lightweight React (Vite)', 'Component reuse; Jacob knows React'],
    ['Backend', 'Node.js + Express 4.x', 'Jacob\'s primary backend stack'],
    ['Auth', 'bcrypt + jsonwebtoken', 'Standard, secure, no external dependency'],
    ['Database', 'PostgreSQL (Fly.io managed)', 'Reliable, scalable; Fly.io native'],
    ['ORM / Query', 'pg (node-postgres)', 'Lightweight; no ORM overhead'],
    ['Peptide Data Store', 'JSON files (Phase 1–2) → PostgreSQL (Phase 3)', 'JSON first for speed; migrate to DB'],
    ['Fonts', 'Inter + JetBrains Mono (Google Fonts)', 'Scientific, modern'],
    ['Icons', 'Lucide Icons (CDN)', 'Clean, consistent'],
    ['Hosting', 'Fly.io (app + DB)', 'Jacob has existing account'],
    ['Source Control', 'GitHub', 'Jacob has existing account'],
    ['CI/CD', 'GitHub Actions → flyctl deploy', 'Auto-deploy on main merge'],
    ['Container', 'Docker (node:20-alpine)', 'Fly.io native'],
]
add_table(['Layer', 'Technology', 'Rationale'], tech_rows)

doc.add_page_break()

# ── 9. File Structure ────────────────────────────────────────────────────────
add_heading('9. File & Directory Structure', 1)

structure_lines = [
    'peptide-vault/',
    '├── index.html                   # Phase 1 prototype (complete, self-contained)',
    '├── package.json',
    '├── Dockerfile',
    '├── fly.toml',
    '├── .github/',
    '│   └── workflows/',
    '│       └── deploy.yml           # GitHub Actions → fly deploy',
    '├── data/',
    '│   ├── peptides.json            # Master peptide dataset (41+ entries)',
    '│   └── scraped/',
    '│       ├── peptpedia_scrape.txt # Raw peptpedia scrape',
    '│       └── sources/             # Per-peptide supplemental source files',
    '├── public/',
    '│   ├── styles/',
    '│   │   ├── main.css',
    '│   │   ├── cards.css',
    '│   │   ├── filters.css',
    '│   │   └── auth.css',
    '│   ├── scripts/',
    '│   │   ├── app.js              # Main app logic',
    '│   │   ├── filter.js           # Filter/sort engine',
    '│   │   ├── cards.js            # Card render + expand',
    '│   │   ├── auth.js             # Auth modal + JWT',
    '│   │   └── bookmarks.js        # Bookmark state + sync',
    '│   └── assets/',
    '│       └── logo.svg',
    '└── src/',
    '    ├── server.js               # Express app entry',
    '    ├── routes/',
    '    │   ├── peptides.js         # GET /api/peptides, /api/peptides/:id',
    '    │   ├── auth.js             # POST /api/auth/register, /login',
    '    │   └── bookmarks.js        # GET/POST/DELETE /api/bookmarks',
    '    ├── middleware/',
    '    │   └── auth.js             # JWT verification middleware',
    '    └── db/',
    '        ├── db.js               # pg Pool setup',
    '        └── schema.sql          # CREATE TABLE statements',
]
for line in structure_lines:
    add_code(line)

doc.add_page_break()

# ── 10. Development Phases ───────────────────────────────────────────────────
add_heading('10. Development Phases & Timeline', 1)

phase_rows = [
    ['Phase 1', 'HTML Prototype', '1–2 days',
     'index.html with all 41 peptides in JSON; full filter/sort/search/expand UI; mock auth via localStorage; complete design system applied'],
    ['Phase 2', 'Data Pipeline', '1 day',
     'Scrape remaining 32 peptides from peptpedia; augment with PubMed/PubChem; compile data/peptides.json; write per-peptide source TXTs'],
    ['Phase 3', 'Backend API', '1–2 days',
     'Express server; GET /api/peptides; GET /api/peptides/:id; serve static files; port prototype to use API'],
    ['Phase 4', 'Auth + Bookmarks', '1–2 days',
     'User registration + login (bcrypt + JWT); bookmark endpoints; PostgreSQL schema; connect frontend auth modals'],
    ['Phase 5', 'Deploy', '1 day',
     'Dockerfile; fly.toml; GitHub Actions workflow; fly deploy; DNS/HTTPS; smoke test all features'],
    ['Phase 6', 'Polish + QA', '1 day',
     'Responsive design QA; all citations hyperlinked; disclaimer text; loading states; error handling; SEO meta tags'],
]
add_table(['Phase', 'Name', 'Est. Time', 'Deliverables'], phase_rows)

doc.add_page_break()

# ── 11. Deployment ───────────────────────────────────────────────────────────
add_heading('11. Deployment (GitHub + Fly.io)', 1)

add_heading('GitHub Setup', 2, color=(30, 120, 180))
github_steps = [
    'Create repo: github.com/jacobbudnitz/peptide-vault (or similar)',
    'Branch strategy: main (production) + dev (development)',
    'GitHub Actions: .github/workflows/deploy.yml triggers on push to main',
    'Secrets: FLY_API_TOKEN stored as GitHub repo secret',
    'Optional: Dependabot for dependency updates',
]
for s in github_steps:
    add_bullet(s)

add_heading('Fly.io Setup', 2, color=(30, 120, 180))
fly_steps = [
    'fly launch --name peptide-vault (creates fly.toml)',
    'fly postgres create --name peptide-vault-db',
    'fly postgres attach peptide-vault-db',
    'fly secrets set JWT_SECRET=<random-64-char-string>',
    'fly deploy (initial)',
    'fly scale count 1 --region ord (or nearest region)',
    'Auto-HTTPS via Fly.io certificates (no config needed)',
]
for s in fly_steps:
    add_bullet(s)

add_heading('GitHub Actions Workflow (deploy.yml)', 2, color=(30, 120, 180))
workflow_lines = [
    'name: Deploy to Fly.io',
    'on:',
    '  push:',
    '    branches: [main]',
    'jobs:',
    '  deploy:',
    '    runs-on: ubuntu-latest',
    '    steps:',
    '      - uses: actions/checkout@v4',
    '      - uses: superfly/flyctl-actions/setup-flyctl@master',
    '      - run: flyctl deploy --remote-only',
    '        env:',
    '          FLY_API_TOKEN: ${{ secrets.FLY_API_TOKEN }}',
]
for line in workflow_lines:
    add_code(line)

doc.add_page_break()

# ── 12. Future Enhancements ──────────────────────────────────────────────────
add_heading('12. Future Enhancements', 1)

future = [
    'Peptide comparison tool — side-by-side of 2–4 compounds across all fields',
    'Consensus MCP integration — pull live peer-reviewed summaries when available',
    'PubMed live search widget — latest papers for any peptide (NCBI E-utilities)',
    'Interactive sequence viewer — amino acid properties color-coded on sequence',
    'SMILES/3D structure viewer — embed PubChem 3D structure iframe',
    'Community notes — annotated user comments (gated, moderated)',
    'Email newsletter — new peptide additions, research updates',
    'Export — download peptide data as PDF or CSV',
    'Mobile app — React Native or PWA with offline support',
    'Admin dashboard — add/edit peptides via UI without code',
]
for f in future:
    add_bullet(f)

doc.add_page_break()

# ── 13. Open Questions ───────────────────────────────────────────────────────
add_heading('13. Open Questions', 1)

questions = [
    'Site name: "Peptide Vault" or something else? (PeptideAtlas is taken by proteomics)',
    'Consensus MCP: Can you re-check connection? If yes, we use it at data scrape phase.',
    'Domain: Do you want a custom domain (e.g., peptidevault.io) or use fly.dev subdomain?',
    'Peptide count: Start with 41 (peptpedia list) or add additional compounds immediately?',
    'Auth: Social login (Google OAuth) in addition to email/password?',
    'Data freshness: How often should peptide data be refreshed? Manual vs automated re-scrape?',
    'Disclaimer / legal: Do you want a formal medical disclaimer page?',
    'Branding: Should the logo/color scheme tie to your other sites (organism-logger, jacobbudnitz.com)?',
]
for q in questions:
    add_bullet(q)

# ── Footer ───────────────────────────────────────────────────────────────────
doc.add_page_break()
footer_text = doc.add_paragraph(
    f'Peptide Vault Implementation Plan  |  Prepared April 20, 2026  |  Jacob Budnitz\n'
    f'Generated by Claude Sonnet 4.6  |  Version 1.0'
)
footer_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_text.runs[0].font.size = Pt(9)
footer_text.runs[0].font.color.rgb = RGBColor(130, 130, 130)

# ── Save ─────────────────────────────────────────────────────────────────────
output_path = '/Users/jacobbudnitz/peptide-vault/Peptide_Vault_Plan.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
